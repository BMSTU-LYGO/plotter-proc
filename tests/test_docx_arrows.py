from pathlib import Path

from docx import Document
from docx.oxml import parse_xml

from plotter_processor.document_models import SourceArrowElement
from plotter_processor.structured_document_reader import read_structured_document


def test_docx_vml_arrow_direction_is_preserved(tmp_path: Path) -> None:
    document = read_structured_document(Path("tests/fixtures/update_7/lines_tables/arrows.docx"), assets_dir=tmp_path)
    arrow = next(item for item in document.elements if isinstance(item, SourceArrowElement))
    assert not arrow.head_at_start and arrow.head_at_end
    assert arrow.points[0].x_mm < arrow.points[-1].x_mm


def test_one_vml_pict_imports_every_arrow_in_source_order(tmp_path: Path) -> None:
    source = tmp_path / "three-arrows.docx"
    document = Document()
    run = document.add_paragraph().add_run()
    run._r.append(parse_xml("""
        <w:pict
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml">
          <v:line from="20,18" to="150,18" strokecolor="black" strokeweight="1.5pt">
            <v:stroke endarrow="open"/>
          </v:line>
          <v:line from="20,47" to="235,47" strokecolor="#112233" strokeweight="2pt">
            <v:stroke startarrow="open" endarrow="open"/>
          </v:line>
          <v:line from="20,78" to="360,78" strokecolor="black" strokeweight="1pt">
            <v:stroke endarrow="classic"/>
          </v:line>
        </w:pict>
    """))
    document.save(source)

    parsed = read_structured_document(source, assets_dir=tmp_path / "assets")
    arrows = [item for item in parsed.elements if isinstance(item, SourceArrowElement)]

    assert len(arrows) == 3
    assert [arrow.source_order for arrow in arrows] == sorted(
        arrow.source_order for arrow in arrows
    )
    assert len({arrow.id for arrow in arrows}) == 3
    assert [(arrow.head_at_start, arrow.head_at_end) for arrow in arrows] == [
        (False, True),
        (True, True),
        (False, True),
    ]
    assert [arrow.head_style for arrow in arrows] == ["open", "open", "classic"]
    assert [arrow.stroke_color for arrow in arrows] == ["black", "#112233", "black"]
    assert [round(arrow.line_width_mm or 0, 3) for arrow in arrows] == [0.529, 0.706, 0.353]
    assert all(arrow.points[0].x_mm < arrow.points[-1].x_mm for arrow in arrows)

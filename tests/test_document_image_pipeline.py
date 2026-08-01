import json
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw

from plotter_processor.pipeline import PipelineOptions, run_pipeline


def test_docx_image_reaches_preview_paths_gcode_and_report(
    tmp_path: Path, test_font: Path
) -> None:
    png = tmp_path / "line.png"
    image = Image.new("RGB", (80, 50), "white")
    ImageDraw.Draw(image).line((5, 5, 70, 40), fill="black", width=4)
    image.save(png)
    source = tmp_path / "mixed.docx"
    document = Document()
    document.add_paragraph("Before image")
    document.add_picture(str(png))
    document.add_paragraph("After image")
    document.save(source)
    output = tmp_path / "build"

    result = run_pipeline(PipelineOptions(
        source, test_font, "A5", "normal", Path("configs/layout.yaml"),
        Path("configs/machine.yaml"), output, images="centerline", image_debug=True,
    ))

    assert result.status == "ok", result.error
    paths = json.loads((output / "paths.json").read_text(encoding="utf-8"))
    assert any(stroke["element_type"] == "raster-image" for stroke in paths["strokes"])
    assert "raster-image" in (output / "plotter-preview.svg").read_text(encoding="utf-8")
    gcode = (output / "output.gcode").read_text(encoding="utf-8")
    assert all(command not in gcode for command in ("M104", "M109", "M140", "M190", "G28"))
    assert not any(
        token.startswith("E") and token[1:].replace(".", "", 1).lstrip("-").isdigit()
        for line in gcode.splitlines()
        for token in line.split()
    )
    report = json.loads((output / "report.json").read_text(encoding="utf-8"))
    assert report["document_import"]["raster_images_vectorized"] == 1
    assert (output / "document-structure.json").is_file()

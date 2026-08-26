from pathlib import Path

from docx import Document
from docx.oxml import parse_xml

from plotter_processor.document_models import SourceTextElement, SourceVectorElement
from plotter_processor.structured_document_reader import read_structured_document


def test_grouped_vml_flowchart_shapes_and_text_are_native_vectors(tmp_path: Path) -> None:
    source = tmp_path / "shapes.docx"
    document = Document()
    run = document.add_paragraph().add_run()
    run._r.append(parse_xml("""
        <w:pict
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml">
          <v:group>
            <v:rect style="position:absolute;left:10pt;top:20pt;width:80pt;height:40pt"
                    fillcolor="none" strokecolor="black">
              <v:textbox><w:txbxContent><w:p><w:r><w:t>Process</w:t></w:r></w:p></w:txbxContent></v:textbox>
            </v:rect>
            <v:roundrect style="position:absolute;left:110pt;top:20pt;width:80pt;height:40pt"
                         fillcolor="none" strokecolor="black"/>
            <v:oval style="position:absolute;left:210pt;top:20pt;width:60pt;height:40pt"
                    fillcolor="none" strokecolor="black"/>
          </v:group>
        </w:pict>
    """))
    document.save(source)

    parsed = read_structured_document(source, assets_dir=tmp_path / "assets")
    vectors = [item for item in parsed.elements if isinstance(item, SourceVectorElement)]
    texts = [item for item in parsed.elements if isinstance(item, SourceTextElement)]

    assert len(vectors) == 3
    assert [item.strokes[0].semantic_role for item in vectors] == [
        "docx-rectangle",
        "docx-rounded-rectangle",
        "docx-ellipse",
    ]
    assert all(item.strokes[0].closed for item in vectors)
    assert all(item.bbox is not None and item.bbox.width > 0 for item in vectors)
    assert any("Process" in paragraph for item in texts for paragraph in item.paragraphs)

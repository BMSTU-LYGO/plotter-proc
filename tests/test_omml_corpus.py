from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml

from plotter_processor.document_models import SourceMathElement
from plotter_processor.latex_renderer import MathTextRenderer
from plotter_processor.omml_parser import parse_omml
from plotter_processor.structured_document_reader import read_structured_document

CORPUS = Path("tests/fixtures/update_18/omml_corpus.xml")


def _corpus_equations() -> list[object]:
    root = parse_xml(CORPUS.read_text(encoding="utf-8"))
    return list(root.iterchildren())


def test_native_omml_corpus_has_semantic_models_and_vector_geometry() -> None:
    parsed = [parse_omml(element) for element in _corpus_equations()]

    assert len(parsed) == 4
    assert all(item.model is not None for item in parsed)
    assert not [warning for item in parsed for warning in item.warnings]
    assert r"\sin x" in parsed[0].expression
    assert r"\frac{a}{b}" in parsed[0].expression
    assert parsed[1].expression == r"\begin{bmatrix}1&2\\3&4\end{bmatrix}"
    assert parsed[2].expression == r"\begin{aligned}x+y=10\\2x-y=5\end{aligned}"
    assert r"\iiint_{0}^{∞}f" in parsed[3].expression

    rendered = [MathTextRenderer(source_kind="omml").render(item.model, 5.0) for item in parsed]
    assert all(item.quality["render_path"] == "vector-first" for item in rendered)
    assert all(item.quality["quality_failures"] == () for item in rendered)
    semantic_kinds = {
        node.kind
        for item in parsed
        for node in _walk(item.model.root)  # type: ignore[union-attr]
    }
    assert {
        "fraction", "superscript", "function", "delimiter", "matrix", "row",
        "equation-array", "accent", "limit", "n-ary", "group",
    } <= semantic_kinds


def test_omml_corpus_survives_native_docx_reading_order(tmp_path: Path) -> None:
    source = tmp_path / "omml-university.docx"
    document = Document()
    document.add_paragraph("Before equations")
    for equation in _corpus_equations():
        paragraph = document.add_paragraph()
        paragraph._p.append(deepcopy(equation))
    document.add_paragraph("After equations")
    document.save(source)

    parsed = read_structured_document(source, assets_dir=tmp_path / "assets")
    math = [element for element in parsed.elements if isinstance(element, SourceMathElement)]

    assert len(math) == 4
    assert all(element.model is not None for element in math)
    assert [element.source_syntax for element in math] == ["omml"] * 4


def _walk(node):
    yield node
    for child in node.children:
        yield from _walk(child)

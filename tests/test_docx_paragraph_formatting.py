from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from plotter_processor.document_models import SourceTableElement, SourceTextElement
from plotter_processor.structured_document_reader import read_structured_document
from tests.test_paragraph_layout import _layout


def _paragraphs(path: Path, tmp_path: Path):
    result = read_structured_document(path, assets_dir=tmp_path / "assets")
    return [
        item.styled_paragraphs[0]
        for item in result.elements
        if isinstance(item, SourceTextElement) and item.styled_paragraphs
    ]


def test_docx_direct_alignment_indents_spacing_and_tabs(tmp_path: Path) -> None:
    document = Document()
    center = document.add_paragraph("center")
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right = document.add_paragraph("right")
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    justify = document.add_paragraph("justify words")
    justify.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    formatted = document.add_paragraph("formatted\tvalue")
    formatted.paragraph_format.left_indent = Mm(8)
    formatted.paragraph_format.right_indent = Mm(6)
    formatted.paragraph_format.first_line_indent = Mm(10)
    formatted.paragraph_format.space_before = Pt(12)
    formatted.paragraph_format.space_after = Pt(6)
    formatted.paragraph_format.line_spacing = 1.5
    formatted.paragraph_format.tab_stops.add_tab_stop(Mm(40))
    hanging = document.add_paragraph("1. hanging")
    hanging.paragraph_format.left_indent = Mm(12)
    hanging.paragraph_format.first_line_indent = Mm(-5)
    path = tmp_path / "formatting.docx"
    document.save(path)

    paragraphs = _paragraphs(path, tmp_path)

    assert [item.alignment for item in paragraphs[:3]] == ["center", "right", "justify"]
    assert abs(paragraphs[3].left_indent_mm - 8) < 0.05
    assert abs(paragraphs[3].right_indent_mm - 6) < 0.05
    assert abs(paragraphs[3].first_line_indent_mm - 10) < 0.05
    assert abs(paragraphs[3].space_before_mm - 12 * 25.4 / 72) < 0.05
    assert paragraphs[3].line_spacing == 1.5
    assert abs(paragraphs[3].tab_stops_mm[0] - 40) < 0.05
    assert paragraphs[4].hanging_indent_mm and paragraphs[4].first_line_indent_mm is None


def test_style_inheritance_direct_override_and_semantic_roles(tmp_path: Path) -> None:
    document = Document()
    base = document.styles.add_style("Base Body", WD_STYLE_TYPE.PARAGRAPH)
    base.paragraph_format.left_indent = Mm(7)
    derived = document.styles.add_style("CustomHeading", WD_STYLE_TYPE.PARAGRAPH)
    derived.base_style = base
    derived.paragraph_format.space_after = Pt(9)
    paragraph = document.add_paragraph("Inherited", style=derived)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph("Title text", style="Title")
    document.add_paragraph("Heading text", style="Heading 1")
    path = tmp_path / "styles.docx"
    document.save(path)

    paragraphs = _paragraphs(path, tmp_path)

    assert abs(paragraphs[0].left_indent_mm - 7) < 0.05
    assert paragraphs[0].space_after_mm is not None
    assert paragraphs[0].alignment == "right"
    assert paragraphs[0].semantic_role == "body"
    assert paragraphs[1].semantic_role == "title"
    assert paragraphs[2].semantic_role == "heading_1"


def test_tab_kinds_and_table_cell_format_are_preserved(
    tmp_path: Path, test_font: Path
) -> None:
    document = Document()
    for kind in ("left", "center", "right", "decimal"):
        token = "12,34" if kind == "decimal" else "value"
        paragraph = document.add_paragraph(f"{kind}\t{token}")
        ppr = paragraph._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), kind)
        tab.set(qn("w:pos"), "1440")
        tabs.append(tab)
        ppr.append(tabs)
    cell_paragraph = document.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
    cell_paragraph.text = "Cell"
    cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = tmp_path / "table.docx"
    document.save(path)

    result = read_structured_document(path, assets_dir=tmp_path / "assets")
    table = next(item for item in result.elements if isinstance(item, SourceTableElement))

    paragraphs = _paragraphs(path, tmp_path)
    assert [item.tab_stops[0].alignment for item in paragraphs[:4]] == [
        "left", "center", "right", "decimal"
    ]
    assert not any(warning.startswith("docx_tab_stop_approximated") for warning in result.warnings)
    assert table.cells[0].paragraphs[0].alignment == "center"
    expected = 10 + 25.4
    for kind, paragraph in zip(("left", "center", "right", "decimal"), paragraphs):
        layout = _layout(paragraph, test_font, right=100)
        glyphs = [glyph for glyph in layout.lines[0].glyphs if glyph.word_index == 1]
        left = min(glyph.x_mm for glyph in glyphs)
        right = max(glyph.x_mm + glyph.advance_mm for glyph in glyphs)
        if kind == "left":
            anchor = left
        elif kind == "center":
            anchor = (left + right) / 2
        elif kind == "right":
            anchor = right
        else:
            anchor = next(glyph.x_mm for glyph in glyphs if glyph.char == ",")
        assert abs(anchor - expected) < 0.01

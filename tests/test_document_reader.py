from pathlib import Path

import pymupdf
import pytest
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

from plotter_processor.document_reader import PDF_TEXT_LAYER_ERROR, read_document


def test_reads_docx_paragraphs_and_preserves_empty_paragraphs(tmp_path: Path) -> None:
    source = tmp_path / "input.docx"
    document = Document()
    document.add_paragraph("Первый абзац")
    document.add_paragraph("")
    document.add_paragraph("Second paragraph")
    document.save(source)

    result = read_document(source)

    assert result.paragraphs == ["Первый абзац", "", "Second paragraph"]
    assert result.source_path == source
    assert result.warnings == []


def test_docx_extract_preserves_paragraph_table_math_reading_order(tmp_path: Path) -> None:
    source = tmp_path / "structured.docx"
    document = Document()
    document.add_paragraph("Before table")
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Merged header"
    table.cell(0, 2).text = "Header 2"
    table.cell(1, 0).text = "Row 2 cell 1"
    table.cell(1, 1).text = "Row 2 cell 2"
    table.cell(1, 2).text = "Row 2 cell 3"
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_properties.append(OxmlElement("w:tblHeader"))
    math_paragraph = document.add_paragraph()
    math_paragraph._p.append(parse_xml(
        f"<m:oMathPara {nsdecls('m')}><m:oMath><m:f>"
        "<m:num><m:r><m:t>x</m:t></m:r></m:num>"
        "<m:den><m:r><m:t>y</m:t></m:r></m:den>"
        "</m:f></m:oMath></m:oMathPara>"
    ))
    document.add_paragraph("After math")
    document.save(source)

    result = read_document(source)

    assert result.paragraphs == [
        "Before table",
        "Merged header",
        "Header 2",
        "Row 2 cell 1",
        "Row 2 cell 2",
        "Row 2 cell 3",
        r"\frac{x}{y}",
        "After math",
    ]


def test_reads_pdf_text_layer_in_page_order(tmp_path: Path) -> None:
    source = tmp_path / "input.pdf"
    document = pymupdf.open()
    first_page = document.new_page()
    first_page.insert_text((72, 72), "First page text")
    second_page = document.new_page()
    second_page.insert_text((72, 72), "Second page text")
    document.save(source)
    document.close()

    result = read_document(source)

    extracted = "\n".join(result.paragraphs)
    assert "First page text" in extracted
    assert "Second page text" in extracted
    assert extracted.index("First page text") < extracted.index("Second page text")


def test_rejects_pdf_without_usable_text_layer(tmp_path: Path) -> None:
    source = tmp_path / "empty.pdf"
    document = pymupdf.open()
    document.new_page()
    document.save(source)
    document.close()

    with pytest.raises(ValueError, match="OCR is not supported"):
        read_document(source)

    assert PDF_TEXT_LAYER_ERROR.endswith("MVP.")


def test_rejects_unsupported_extension(tmp_path: Path) -> None:
    source = tmp_path / "input.rtf"
    source.write_text("plain text", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported input format"):
        read_document(source)


def test_reads_utf8_bom_txt_and_preserves_lines(tmp_path: Path) -> None:
    source = tmp_path / "input.txt"
    source.write_text("\ufeffПривет\n\nЁжик", encoding="utf-8")

    result = read_document(source)

    assert result.paragraphs == ["Привет", "", "Ёжик"]


def test_rejects_empty_txt(tmp_path: Path) -> None:
    source = tmp_path / "empty.txt"
    source.write_text("  \n", encoding="utf-8")

    with pytest.raises(ValueError, match="no usable text"):
        read_document(source)

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageDraw

from plotter_processor.document_models import (
    SourceRasterImageElement,
    SourceTableElement,
    SourceTextElement,
)
from plotter_processor.structured_document_reader import read_structured_document


def _png(path: Path) -> None:
    image = Image.new("RGBA", (40, 30), (255, 255, 255, 0))
    ImageDraw.Draw(image).line((2, 2, 36, 25), fill=(0, 0, 0, 255), width=3)
    image.save(path)


def test_docx_inline_image_preserves_run_order_and_table_image(tmp_path: Path) -> None:
    png = tmp_path / "line.png"
    _png(png)
    source = tmp_path / "mixed.docx"
    docx = Document()
    paragraph = docx.add_paragraph()
    paragraph.add_run("before")
    paragraph.add_run().add_picture(str(png))
    paragraph.add_run("after")
    cell = docx.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run().add_picture(str(png))
    docx.save(source)

    result = read_structured_document(source, assets_dir=tmp_path / "assets")
    elements = result.pages[0].elements

    assert isinstance(elements[0], SourceTextElement)
    assert isinstance(elements[1], SourceRasterImageElement)
    assert isinstance(elements[2], SourceTextElement)
    assert elements[0].paragraphs == ("before",)
    assert elements[2].paragraphs == ("after",)
    assert sum(isinstance(item, SourceRasterImageElement) for item in elements) == 1
    assert any(isinstance(item, SourceTableElement) for item in elements)
    assert "docx_table_embedded_object_not_supported" in result.warnings


def test_docx_anchor_is_not_ignored(tmp_path: Path) -> None:
    png = tmp_path / "line.png"
    _png(png)
    source = tmp_path / "anchor.docx"
    docx = Document()
    shape = docx.add_paragraph().add_run().add_picture(str(png))
    shape._inline.tag = qn("wp:anchor")
    docx.save(source)

    result = read_structured_document(source, assets_dir=tmp_path / "assets")

    assert any(isinstance(item, SourceRasterImageElement) for item in result.elements)
    image = next(item for item in result.elements if isinstance(item, SourceRasterImageElement))
    assert image.anchor_type == "anchored"
    assert image.wrap_mode == "square"
    assert "floating_image_reflowed" not in result.warnings

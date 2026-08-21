from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from PIL import Image, ImageDraw

OUTPUT = Path("tests/fixtures/layout/paragraph_formatting_demo.docx")


def main() -> None:
    document = Document()
    title = document.add_paragraph("ЛАБОРАТОРНАЯ РАБОТА №1", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Форматирование абзацев", style="Heading 1")
    first = document.add_paragraph(
        "Первый абзац начинается с красной строки и продолжается " * 4
    )
    first.paragraph_format.first_line_indent = Mm(10)
    left = document.add_paragraph("Абзац с левым отступом. " * 5)
    left.paragraph_format.left_indent = Mm(12)
    right_indent = document.add_paragraph("Абзац с правым отступом. " * 5)
    right_indent.paragraph_format.right_indent = Mm(15)
    centered = document.add_paragraph("Центрированный абзац")
    centered.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right = document.add_paragraph("Абзац по правому краю")
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    justified = document.add_paragraph("Выровненный по ширине длинный абзац. " * 10)
    justified.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    tabs = document.add_paragraph("Термин:\tЗначение\nСкорость:\t1000 мм/мин")
    tabs.paragraph_format.tab_stops.add_tab_stop(Mm(55))
    formula = document.add_paragraph("$$E = mc^2$$")
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER

    image_path = OUTPUT.parent / ".paragraph-demo-image.png"
    image = Image.new("RGB", (480, 160), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((10, 10, 470, 150), outline="black", width=5)
    draw.line((30, 130, 230, 35, 450, 130), fill="black", width=6)
    image.save(image_path)
    document.add_paragraph().add_run().add_picture(str(image_path), width=Mm(90))
    document.add_page_break()
    document.add_paragraph("Вторая страница", style="Heading 1")
    document.add_paragraph("Формат абзаца должен сохраниться после page break. " * 8)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    image_path.unlink()
    print(OUTPUT)


if __name__ == "__main__":
    main()

"""Generate deterministic local fixtures for UPD_Plotter_7.

The files intentionally cover unsupported baseline behaviour as well as the
candidate scenarios.  No network resources or installed office applications
are used.
"""

from __future__ import annotations

import argparse
import io
import zipfile
from datetime import UTC, datetime
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Mm
from PIL import Image, ImageDraw

FIXED_ZIP_TIME = (2020, 1, 1, 0, 0, 0)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output", type=Path, default=Path("tests/fixtures/update_7")
    )
    parser.add_argument(
        "--examples-output", type=Path, default=Path("examples/update_7/block_1")
    )
    parser.add_argument(
        "--block-2-examples-output",
        type=Path,
        default=Path("examples/update_7/block_2"),
    )
    args = parser.parse_args()
    root = args.output
    latex = root / "latex"
    images = root / "images"
    lines = root / "lines_tables"
    for directory in (latex, images, lines):
        directory.mkdir(parents=True, exist_ok=True)

    _write_text_fixtures(latex)
    _write_latex_docx(latex / "latex_in_docx.docx")
    _write_omml_docx(latex / "omml_basic.docx")
    _write_formula_pdf(latex / "pdf_formula_text.pdf", vector_fraction=False)
    _write_formula_pdf(latex / "pdf_formula_vector.pdf", vector_fraction=True)

    png = _sample_image()
    _write_image_docx(images / "image_left_wrap.docx", png, "left", "square")
    _write_image_docx(images / "image_right_wrap.docx", png, "right", "square")
    _write_image_docx(images / "image_top_bottom.docx", png, "center", "top_bottom")
    _write_image_docx(images / "image_absolute.docx", png, "right", "none")
    _write_image_pdf(images / "image_left_wrap.pdf", png, "left")
    _write_image_pdf(images / "image_right_wrap.pdf", png, "right")
    _write_image_pdf(images / "image_preserve_position.pdf", png, "right")
    _write_overlap_pdf(images / "image_overlap.pdf", png)

    _write_underlines_docx(lines / "underline_runs.docx")
    _write_underline_pdf(lines / "underline_pdf.pdf")
    _write_arrows_docx(lines / "arrows.docx")
    _write_arrows_pdf(lines / "arrows.pdf")
    _write_table_docx(lines / "simple_table.docx", rows=3, columns=3)
    _write_merged_table_docx(lines / "merged_cells.docx")
    _write_table_docx(lines / "multipage_table.docx", rows=36, columns=3)
    _write_table_pdf(lines / "simple_table.pdf")
    _write_table_with_underlines(lines / "table_with_underlines.docx")
    _write_table_with_arrows(lines / "table_with_arrows.docx")

    examples = args.examples_output
    examples.mkdir(parents=True, exist_ok=True)
    _write_block_1_example_docx(examples / "semantic-omml.docx")
    _write_block_1_example_pdf(examples / "pdf-visual-math.pdf")

    block_2_examples = args.block_2_examples_output
    block_2_examples.mkdir(parents=True, exist_ok=True)
    _write_image_docx(
        block_2_examples / "image-left-square-wrap.docx", png, "left", "square"
    )
    _write_image_docx(
        block_2_examples / "image-right-square-wrap.docx", png, "right", "square"
    )
    _write_image_pdf(block_2_examples / "pdf-image-right.pdf", png, "right")
    return 0


def _write_text_fixtures(directory: Path) -> None:
    values = {
        "latex_inline.txt": "Pythagoras: $x^2 + y^2 = z^2$.\n",
        "latex_block.txt": "Block formula:\n\n$$\\frac{a+b}{c-d}$$\n",
        "latex_complex.txt": (
            "Inline $x_i + \\alpha \\leq \\beta$.\n\n"
            "$$\\int_0^1 f(x)\\,dx = \\frac{1}{3}$$\n\n"
            "$$\\sum_{i=1}^{n} i + \\sqrt{x^2 + 1}$$\n"
        ),
        "latex_multipage.txt": "Line\n" * 42 + "$$\\frac{a}{b}$$\n",
    }
    for name, value in values.items():
        (directory / name).write_text(value, encoding="utf-8")


def _base_docx() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(148)
    section.page_height = Mm(210)
    section.left_margin = section.right_margin = Mm(12)
    section.top_margin = section.bottom_margin = Mm(12)
    document.core_properties.title = "UPD Plotter 7 fixture"
    document.core_properties.author = "plotter-processor"
    fixed_time = datetime(2020, 1, 1, tzinfo=UTC)
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time
    return document


def _save_docx(document: Document, path: Path) -> None:
    temporary = io.BytesIO()
    document.save(temporary)
    source = zipfile.ZipFile(io.BytesIO(temporary.getvalue()))
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, source.read(name))
    path.write_bytes(output.getvalue())


def _write_latex_docx(path: Path) -> None:
    document = _base_docx()
    document.add_paragraph("Word text with $x^2 + \\sqrt{x}$ and normal text.")
    document.add_paragraph("$$\\sum_{i=1}^{n} i$$")
    _save_docx(document, path)


def _write_omml_docx(path: Path) -> None:
    document = _base_docx()
    document.add_paragraph("Basic Word equation:")
    paragraph = document.add_paragraph()
    math = parse_xml(
        f"<m:oMathPara {nsdecls('m')}>"
        "<m:oMath><m:f><m:num><m:r><m:t>a+b</m:t></m:r></m:num>"
        "<m:den><m:r><m:t>c</m:t></m:r></m:den></m:f>"
        "<m:r><m:t>=</m:t></m:r>"
        "<m:sRad><m:deg/><m:e><m:r><m:t>x</m:t></m:r></m:e></m:sRad>"
        "</m:oMath></m:oMathPara>"
    )
    paragraph._p.append(math)
    _save_docx(document, path)


def _write_block_1_example_docx(path: Path) -> None:
    document = _base_docx()
    document.core_properties.title = "UPD Plotter 7 block 1 example"
    document.add_heading("Block 1: semantic and Word formulas", level=1)
    document.add_paragraph(
        "Ordinary centerline text with an inline formula: "
        "$x_i + \\alpha \\leq \\beta$."
    )
    document.add_paragraph("A block MathText formula follows:")
    formula = document.add_paragraph("$$\\int_0^1 f(x)\\,dx = \\frac{1}{3}$$")
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("A native Word OMML equation follows:")
    paragraph = document.add_paragraph()
    math = parse_xml(
        f"<m:oMathPara {nsdecls('m')}>"
        "<m:oMath><m:f><m:num><m:r><m:t>a+b</m:t></m:r></m:num>"
        "<m:den><m:r><m:t>c</m:t></m:r></m:den></m:f>"
        "<m:r><m:t>=</m:t></m:r>"
        "<m:sRad><m:deg/><m:e><m:sSup><m:e><m:r><m:t>x</m:t></m:r></m:e>"
        "<m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>"
        "</m:e></m:sRad></m:oMath></m:oMathPara>"
    )
    paragraph._p.append(math)
    document.add_paragraph(
        "The three formulas above must use single centerlines, never doubled outlines."
    )
    _save_docx(document, path)


def _new_pdf() -> tuple[fitz.Document, fitz.Page]:
    document = fitz.open()
    return document, document.new_page(width=419.528, height=595.276)


def _save_pdf(document: fitz.Document, path: Path) -> None:
    metadata = dict(document.metadata)
    metadata.update({
        "producer": "plotter-processor fixture generator",
        "creator": "plotter-processor fixture generator",
        "creationDate": "D:20200101000000Z",
        "modDate": "D:20200101000000Z",
    })
    document.set_metadata(metadata)
    document.save(path, garbage=4, deflate=True, clean=True, no_new_id=True)
    document.close()


def _write_formula_pdf(path: Path, *, vector_fraction: bool) -> None:
    document, page = _new_pdf()
    page.insert_text((45, 55), "Formula region", fontsize=12)
    if vector_fraction:
        page.insert_text((130, 92), "a+b", fontsize=13)
        page.draw_line((126, 98), (166, 98), width=1)
        page.insert_text((137, 115), "c", fontsize=13)
        page.insert_text((174, 102), "= x^2", fontsize=15)
    else:
        page.insert_text((85, 100), "x^2 + y^2 = z^2", fontsize=18)
        page.insert_text((85, 132), "sum i=1..n", fontsize=14)
    document.set_metadata({"title": "UPD Plotter 7 formula fixture"})
    _save_pdf(document, path)


def _write_block_1_example_pdf(path: Path) -> None:
    document, page = _new_pdf()
    page.insert_text((45, 45), "Block 1: PDF visual formulas", fontsize=15)
    page.insert_text((45, 70), "High-confidence text formula:", fontsize=10)
    page.insert_text((85, 105), "x^2 + y^2 = z^2", fontsize=18)

    page.insert_text((45, 175), "Formula containing a vector fraction bar:", fontsize=10)
    page.insert_text((130, 215), "a+b", fontsize=14)
    page.draw_line((126, 221), (168, 221), width=1)
    page.insert_text((140, 240), "c", fontsize=14)
    page.insert_text((178, 228), "= sqrt(x)", fontsize=16)

    page.insert_text(
        (45, 290),
        "Each detected region is clipped once and replaces its source primitives.",
        fontsize=9,
    )
    document.set_metadata({"title": "UPD Plotter 7 block 1 PDF example"})
    _save_pdf(document, path)


def _sample_image() -> bytes:
    image = Image.new("RGB", (240, 160), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((4, 4, 235, 155), outline="black", width=4)
    draw.ellipse((55, 25, 185, 140), outline="black", width=5)
    draw.line((15, 145, 225, 15), fill="black", width=4)
    output = io.BytesIO()
    image.save(output, format="PNG", optimize=False)
    return output.getvalue()


def _write_image_docx(path: Path, png: bytes, side: str, wrap: str) -> None:
    document = _base_docx()
    document.add_paragraph(f"Anchored image: side={side}, wrap={wrap}.")
    paragraph = document.add_paragraph()
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }[side]
    run = paragraph.add_run()
    inline = run.add_picture(io.BytesIO(png), width=Mm(42))._inline
    _convert_inline_to_anchor(inline, side, wrap)
    document.add_paragraph(
        "Text around the image. This deterministic paragraph is long enough "
        "to exercise flow and placement decisions on the target page. " * 3
    )
    _save_docx(document, path)


def _convert_inline_to_anchor(inline, side: str, wrap: str) -> None:
    inline.tag = qn("wp:anchor")
    for key, value in {
        "distT": "72000",
        "distB": "72000",
        "distL": "72000",
        "distR": "72000",
        "simplePos": "0",
        "relativeHeight": "2",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "0",
    }.items():
        inline.set(key, value)
    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    horizontal = OxmlElement("wp:positionH")
    horizontal.set("relativeFrom", "margin")
    align = OxmlElement("wp:align")
    align.text = side
    horizontal.append(align)
    vertical = OxmlElement("wp:positionV")
    vertical.set("relativeFrom", "paragraph")
    offset = OxmlElement("wp:posOffset")
    offset.text = "0"
    vertical.append(offset)
    wrapping = OxmlElement(
        {"square": "wp:wrapSquare", "top_bottom": "wp:wrapTopAndBottom", "none": "wp:wrapNone"}[wrap]
    )
    if wrap == "square":
        wrapping.set("wrapText", "bothSides")
    inline.insert(0, simple)
    inline.insert(1, horizontal)
    inline.insert(2, vertical)
    inline.insert(3, wrapping)


def _write_image_pdf(path: Path, png: bytes, side: str) -> None:
    document, page = _new_pdf()
    rect = fitz.Rect(30, 80, 150, 160) if side == "left" else fitz.Rect(270, 80, 390, 160)
    page.insert_image(rect, stream=png)
    text_x = 175 if side == "left" else 35
    page.insert_textbox(
        fitz.Rect(text_x, 75, 390 if side == "left" else 250, 210),
        "Text beside a positioned image. " * 10,
        fontsize=10,
    )
    _save_pdf(document, path)


def _write_overlap_pdf(path: Path, png: bytes) -> None:
    document, page = _new_pdf()
    page.insert_textbox(fitz.Rect(40, 60, 360, 180), "Overlapping text region. " * 16, fontsize=11)
    page.insert_image(fitz.Rect(120, 95, 300, 215), stream=png)
    _save_pdf(document, path)


def _write_underlines_docx(path: Path) -> None:
    document = _base_docx()
    paragraph = document.add_paragraph()
    paragraph.add_run("Single underline").underline = True
    paragraph.add_run(" and normal text")
    double = document.add_paragraph().add_run("Double underline")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "double")
    double._r.get_or_add_rPr().append(underline)
    words = document.add_paragraph().add_run("Words only underline")
    words_only = OxmlElement("w:u")
    words_only.set(qn("w:val"), "words")
    words._r.get_or_add_rPr().append(words_only)
    _save_docx(document, path)


def _write_underline_pdf(path: Path) -> None:
    document, page = _new_pdf()
    page.insert_text((55, 85), "Underlined PDF text", fontsize=16)
    page.draw_line((55, 89), (200, 89), width=0.8)
    page.insert_text((55, 125), "Generic line below", fontsize=12)
    page.draw_line((55, 150), (250, 150), width=1.2)
    _save_pdf(document, path)


def _write_arrows_docx(path: Path) -> None:
    document = _base_docx()
    document.add_paragraph("VML arrow connector:")
    run = document.add_paragraph().add_run()
    pict = parse_xml(
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml">'
        '<v:line from="0,0" to="120pt,40pt" strokecolor="black">'
        '<v:stroke endarrow="block"/></v:line></w:pict>'
    )
    run._r.append(pict)
    _save_docx(document, path)


def _write_arrows_pdf(path: Path) -> None:
    document, page = _new_pdf()
    page.draw_line((55, 80), (260, 80), width=1.2)
    page.draw_polyline([(260, 80), (244, 72), (244, 88), (260, 80)], color=(0, 0, 0), fill=(0, 0, 0))
    page.draw_line((280, 140), (90, 210), width=1.1)
    page.draw_polyline([(90, 210), (102, 195), (108, 211)], color=(0, 0, 0), width=1.1)
    _save_pdf(document, path)


def _write_table_docx(path: Path, *, rows: int, columns: int) -> None:
    document = _base_docx()
    table = document.add_table(rows=rows, cols=columns)
    table.style = "Table Grid"
    for row in range(rows):
        for column in range(columns):
            table.cell(row, column).text = f"R{row + 1}C{column + 1}"
    if rows > 10:
        header = table.rows[0]._tr.get_or_add_trPr()
        header.append(OxmlElement("w:tblHeader"))
    _save_docx(document, path)


def _write_merged_table_docx(path: Path) -> None:
    document = _base_docx()
    table = document.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Horizontal merge"
    table.cell(1, 0).merge(table.cell(2, 0)).text = "Vertical merge"
    for row in range(4):
        for column in range(4):
            if not table.cell(row, column).text:
                table.cell(row, column).text = f"{row + 1},{column + 1}"
    _save_docx(document, path)


def _write_table_pdf(path: Path) -> None:
    document, page = _new_pdf()
    left, top, cell_w, cell_h = 55, 70, 90, 35
    for row in range(4):
        page.draw_line((left, top + row * cell_h), (left + 3 * cell_w, top + row * cell_h))
    for column in range(4):
        page.draw_line((left + column * cell_w, top), (left + column * cell_w, top + 3 * cell_h))
    for row in range(3):
        for column in range(3):
            page.insert_text((left + column * cell_w + 8, top + row * cell_h + 22), f"{row + 1},{column + 1}", fontsize=10)
    _save_pdf(document, path)


def _write_table_with_underlines(path: Path) -> None:
    document = _base_docx()
    table = document.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for row in range(3):
        for column in range(3):
            run = table.cell(row, column).paragraphs[0].add_run(f"Cell {row + 1},{column + 1}")
            run.underline = row == 1
    _save_docx(document, path)


def _write_table_with_arrows(path: Path) -> None:
    document = _base_docx()
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Formula $x^2$"
    table.cell(0, 1).text = "Arrow ->"
    table.cell(1, 0).text = "Start"
    table.cell(1, 1).text = "End"
    run = table.cell(1, 0).paragraphs[0].add_run()
    run._r.append(parse_xml(
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml"><v:line from="0,0" to="50pt,0">'
        '<v:stroke endarrow="block"/></v:line></w:pict>'
    ))
    _save_docx(document, path)


if __name__ == "__main__":
    raise SystemExit(main())

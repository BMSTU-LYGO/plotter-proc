"""Generate deterministic mixed-layout DOCX and PDF regression fixtures."""

from __future__ import annotations

import io
import zipfile
from datetime import UTC, datetime
from pathlib import Path

import pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm
from PIL import Image, ImageDraw

OUTPUT_DIR = Path("tests/fixtures/layout")
FIXED_ZIP_TIME = (2020, 1, 1, 0, 0, 0)


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    image = _sample_image()
    _write_docx(OUTPUT_DIR / "mixed_layout_demo.docx", image)
    _write_pdf(OUTPUT_DIR / "mixed_layout_demo.pdf", image)
    return 0


def _sample_image() -> bytes:
    image = Image.new("RGB", (300, 180), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((4, 4, 295, 175), outline="black", width=5)
    draw.ellipse((75, 25, 225, 155), outline="black", width=5)
    draw.line((20, 155, 280, 20), fill="black", width=5)
    output = io.BytesIO()
    image.save(output, format="PNG", optimize=False)
    return output.getvalue()


def _base_docx() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = section.right_margin = Mm(18)
    section.top_margin = section.bottom_margin = Mm(18)
    fixed = datetime(2020, 1, 1, tzinfo=UTC)
    document.core_properties.title = "Mixed layout regression fixture"
    document.core_properties.author = "plotter-processor"
    document.core_properties.created = fixed
    document.core_properties.modified = fixed
    return document


def _write_docx(path: Path, image: bytes) -> None:
    document = _base_docx()
    document.add_heading("Mixed layout regression", level=1)
    document.add_paragraph(
        "Early text must remain at the top and must not be displaced by a future image. "
        "This paragraph intentionally precedes the anchored object in source order."
    )
    document.add_paragraph(
        "Second early paragraph. Its line rhythm should be stable before the image anchor. "
        "The image belongs to a later source element."
    )
    anchor_paragraph = document.add_paragraph()
    anchor_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    inline = anchor_paragraph.add_run().add_picture(io.BytesIO(image), width=Mm(58))._inline
    _convert_to_anchor(inline, relative_to_v="page", y_mm=48.0, side="right")
    document.add_paragraph(
        "Text after the first image wraps beside it and returns to the full width below it. " * 4
    )
    document.add_paragraph("Inline formula: $x^2 + y^2 = r^2$ continues in this line.")
    document.add_paragraph("$$\\int_0^\\infty x^2 e^{-x}\\,dx = 2$$")
    for index in range(1, 23):
        document.add_paragraph(
            f"Flow paragraph {index:02d}. Natural pagination keeps source order and stable "
            "spacing while the mixed document continues across pages."
        )
        if index == 9:
            middle = document.add_paragraph()
            inline_middle = middle.add_run().add_picture(io.BytesIO(image), width=Mm(52))._inline
            _convert_to_anchor(
                inline_middle, relative_to_v="paragraph", y_mm=0.0, side="left"
            )
    _save_docx(document, path)


def _convert_to_anchor(inline, *, relative_to_v: str, y_mm: float, side: str) -> None:
    inline.tag = qn("wp:anchor")
    for key, value in {
        "distT": "72000",
        "distB": "72000",
        "distL": "72000",
        "distR": "72000",
        "simplePos": "0",
        "relativeHeight": "2",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "0",
    }.items():
        inline.set(key, value)
    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    horizontal = OxmlElement("wp:positionH")
    horizontal.set("relativeFrom", "margin")
    align = OxmlElement("wp:align")
    align.text = side
    horizontal.append(align)
    vertical = OxmlElement("wp:positionV")
    vertical.set("relativeFrom", relative_to_v)
    offset = OxmlElement("wp:posOffset")
    offset.text = str(round(y_mm * 36000))
    vertical.append(offset)
    wrapping = OxmlElement("wp:wrapSquare")
    wrapping.set("wrapText", "bothSides")
    inline.insert(0, simple)
    inline.insert(1, horizontal)
    inline.insert(2, vertical)
    inline.insert(3, wrapping)


def _save_docx(document: Document, path: Path) -> None:
    temporary = io.BytesIO()
    document.save(temporary)
    with zipfile.ZipFile(io.BytesIO(temporary.getvalue())) as source:
        output = io.BytesIO()
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as target:
            for name in sorted(source.namelist()):
                info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o600 << 16
                target.writestr(info, source.read(name))
    path.write_bytes(output.getvalue())


def _write_pdf(path: Path, image: bytes) -> None:
    document = pymupdf.open()
    for page_index in range(2):
        page = document.new_page(width=595.276, height=841.89)
        page.insert_text((52, 52), f"Mixed layout regression - page {page_index + 1}", fontsize=15)
        page.insert_textbox(
            pymupdf.Rect(52, 72, 540, 160),
            "Natural text flow before the image. " * 12,
            fontsize=10,
        )
        image_rect = (
            pymupdf.Rect(365, 175, 535, 277)
            if page_index == 0
            else pymupdf.Rect(52, 220, 222, 322)
        )
        page.insert_image(image_rect, stream=image)
        page.insert_textbox(
            pymupdf.Rect(52, 170, 345 if page_index == 0 else 540, 360),
            "Wrapped text remains close to the image and returns to normal width below it. " * 16,
            fontsize=10,
        )
        page.insert_text((100, 390), "Inline: x^2 + y^2 = r^2", fontsize=12)
        page.insert_text((175, 435), "integral(0, infinity) x^2 e^-x dx = 2", fontsize=14)
        page.insert_textbox(
            pymupdf.Rect(52, 470, 540, 710),
            "Additional paragraphs force a genuine multi-page mixed document. " * 28,
            fontsize=10,
        )
    metadata = dict(document.metadata)
    metadata.update({
        "title": "Mixed layout regression fixture",
        "producer": "plotter-processor fixture generator",
        "creator": "plotter-processor fixture generator",
        "creationDate": "D:20200101000000Z",
        "modDate": "D:20200101000000Z",
    })
    document.set_metadata(metadata)
    document.save(path, garbage=4, deflate=True, clean=True, no_new_id=True)
    document.close()


if __name__ == "__main__":
    raise SystemExit(main())

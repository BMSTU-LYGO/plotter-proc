from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pymupdf
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image, ImageDraw

from plotter_processor.latex_renderer import MathTextRenderer

ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "tests" / "fixtures" / "update_18"


def main() -> None:
    TARGET.mkdir(parents=True, exist_ok=True)
    _docx(TARGET / "complex_math_document.docx")
    _pdf(TARGET / "mixed_math_diagram.pdf")
    _math_preview(TARGET / "canonical_math_preview.svg")


def _docx(path: Path) -> None:
    document = Document()
    document.add_heading("UPD 18 math and document corpus", level=1)
    document.add_paragraph(r"Inline formula: $E=mc^2$ and $x_i^2$.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Equation"
    table.cell(0, 1).text = "Meaning"
    table.cell(1, 1).text = "Native OMML fraction"
    table.cell(1, 0)._tc.append(_fraction_omml())
    image = Image.new("RGB", (160, 80), "white")
    ImageDraw.Draw(image).line((10, 70, 150, 10), fill="black", width=4)
    stream = BytesIO()
    image.save(stream, format="PNG")
    stream.seek(0)
    document.add_picture(stream)
    document.add_page_break()
    math = document.add_paragraph()
    math._p.append(_matrix_omml())
    run = document.add_paragraph().add_run()
    run._r.append(parse_xml("""
        <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:v="urn:schemas-microsoft-com:vml">
          <v:rect style="position:absolute;left:20pt;top:30pt;width:100pt;height:45pt"
                  fillcolor="none" strokecolor="black">
            <v:textbox><w:txbxContent><w:p><w:r><w:t>Flow</w:t></w:r></w:p></w:txbxContent></v:textbox>
          </v:rect>
          <v:line from="120pt,52pt" to="190pt,52pt"><v:stroke endarrow="open"/></v:line>
        </w:pict>
    """))
    document.save(path)


def _fraction_omml():
    return parse_xml(
        f"<m:oMath {nsdecls('m')}><m:f>"
        "<m:num><m:r><m:t>x+1</m:t></m:r></m:num>"
        "<m:den><m:r><m:t>x-1</m:t></m:r></m:den>"
        "</m:f></m:oMath>"
    )


def _matrix_omml():
    return parse_xml(
        f"<m:oMathPara {nsdecls('m')}><m:oMath><m:m>"
        "<m:mr><m:e><m:r><m:t>1</m:t></m:r></m:e><m:e><m:r><m:t>2</m:t></m:r></m:e></m:mr>"
        "<m:mr><m:e><m:r><m:t>3</m:t></m:r></m:e><m:e><m:r><m:t>4</m:t></m:r></m:e></m:mr>"
        "</m:m></m:oMath></m:oMathPara>"
    )


def _pdf(path: Path) -> None:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((40, 40), "Ordinary text next to mathematics", fontsize=11)
    page.insert_text((60, 85), "x^2 + y^2 = z^2", fontsize=16)
    page.insert_text((60, 125), "a+b", fontsize=14)
    page.draw_line((58, 130), (100, 130), width=0.7)
    page.insert_text((72, 148), "c", fontsize=14)
    page.draw_rect((180, 70, 260, 115), color=(0, 0, 0))
    page.draw_circle((320, 92), 22, color=(0, 0, 0))
    page.draw_line((260, 92), (298, 92), color=(0, 0, 0))
    document.save(path)
    document.close()


def _math_preview(path: Path) -> None:
    rendered = MathTextRenderer().render(
        r"f(x)=\frac{1}{\sigma\sqrt{2\pi}}e^{-\frac{(x-\mu)^2}{2\sigma^2}}",
        6.0,
    )
    lines = [
        f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {rendered.width_mm} {rendered.height_mm}">',
        '<rect width="100%" height="100%" fill="white"/>',
    ]
    for stroke in rendered.strokes:
        points = " ".join(f"{point.x:.5f},{point.y:.5f}" for point in stroke.points)
        lines.append(f'<polyline points="{points}" fill="none" stroke="black" stroke-width="0.08"/>')
    lines.append("</svg>")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()

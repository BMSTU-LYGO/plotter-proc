from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm
from PIL import Image, ImageDraw

ROOT = Path("tests/fixtures/layout")


def main() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)
    _build(ROOT / "a4_to_a5_layout_demo.docx", 210, 297, full=False)
    _build(ROOT / "a5_to_a4_layout_demo.docx", 148, 210, full=False)
    _build(ROOT / "upd10_full_demo.docx", 210, 297, full=True)
    print("Generated UPD_Plotter_10 layout fixtures")


def _build(path: Path, width: float, height: float, *, full: bool) -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(width)
    section.page_height = Mm(height)
    section.left_margin = section.right_margin = Mm(20 if width > 150 else 12)
    section.top_margin = section.bottom_margin = Mm(18 if height > 250 else 12)
    title = document.add_paragraph("Масштабирование A4 / A5", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Рисунки и таблицы", style="Heading 1")
    intro = document.add_paragraph(
        "Абзац с красной строкой демонстрирует сохранение структуры " * 3
    )
    intro.paragraph_format.first_line_indent = Mm(10)
    small = _image(320, 160, "small right")
    paragraph = document.add_paragraph()
    anchor = paragraph.add_run().add_picture(io.BytesIO(small), width=Mm(38))._inline
    _anchor(anchor, "right", "paragraph", 4, rotation=12)
    document.add_paragraph("Текст рядом с маленьким рисунком. " * 5)
    large = _image(640, 280, "large centered")
    paragraph = document.add_paragraph()
    anchor = paragraph.add_run().add_picture(io.BytesIO(large), width=Mm(115))._inline
    _anchor(anchor, "center", "paragraph", 8)
    for _ in range(3):
        document.add_paragraph("Заполнение перед таблицей. " * 7)
    _table(document)
    for _ in range(7):
        document.add_paragraph("Текст после таблицы для проверки многостраничности. " * 5)
    if full:
        justified = document.add_paragraph("Абзац по ширине для full regression. " * 8)
        justified.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        right = document.add_paragraph("Строка по правому краю")
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        tabs = document.add_paragraph("Параметр:\tЗначение\nСкорость:\t1000")
        tabs.paragraph_format.tab_stops.add_tab_stop(Mm(55))
        document.add_paragraph("Inline formula: $x^2+y^2=z^2$.")
        block = document.add_paragraph("$$E = mc^2$$")
        block.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(path)


def _table(document: Document) -> None:
    table = document.add_table(rows=7, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Mm(24), Mm(72), Mm(36))
    for column, width in zip(table.columns, widths, strict=True):
        column.width = width
        for cell in column.cells:
            cell.width = width
    headers = ("№", "Результат", "Единицы")
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header)
    table.cell(1, 1).merge(table.cell(1, 2)).text = "Горизонтальное объединение"
    table.cell(2, 0).merge(table.cell(3, 0)).text = "2–3"
    for row in range(1, 7):
        table.cell(row, 0).text = table.cell(row, 0).text or str(row)
        if row != 1:
            table.cell(row, 1).text = (
                "Длинный текст в ячейке, который должен переноситься " * (3 if row == 4 else 1)
            )
            table.cell(row, 2).text = "мм"
    table.cell(4, 1).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), "720")
    table.rows[5]._tr.get_or_add_trPr().append(height)


def _image(width: int, height: int, label: str) -> bytes:
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((4, 4, width - 5, height - 5), outline="black", width=4)
    draw.line((15, height - 20, width // 2, 20, width - 15, height - 20), fill="black", width=5)
    draw.text((12, 12), label, fill="black")
    stream = io.BytesIO()
    image.save(stream, format="PNG")
    return stream.getvalue()


def _anchor(inline, side: str, relative_v: str, y_mm: float, *, rotation: float = 0) -> None:
    inline.tag = qn("wp:anchor")
    for key, value in {
        "distT": "72000", "distB": "72000", "distL": "72000", "distR": "72000",
        "simplePos": "0", "relativeHeight": "2", "behindDoc": "0", "locked": "0",
        "layoutInCell": "1", "allowOverlap": "0",
    }.items():
        inline.set(key, value)
    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    horizontal = OxmlElement("wp:positionH")
    horizontal.set("relativeFrom", "margin")
    align = OxmlElement("wp:align")
    align.text = side
    horizontal.append(align)
    vertical = OxmlElement("wp:positionV")
    vertical.set("relativeFrom", relative_v)
    offset = OxmlElement("wp:posOffset")
    offset.text = str(int(y_mm * 36000))
    vertical.append(offset)
    wrapping = OxmlElement("wp:wrapSquare")
    wrapping.set("wrapText", "bothSides")
    inline.insert(0, simple)
    inline.insert(1, horizontal)
    inline.insert(2, vertical)
    inline.insert(3, wrapping)
    if rotation:
        transforms = inline.xpath(".//*[local-name()='xfrm']")
        if transforms:
            transforms[0].set("rot", str(int(rotation * 60000)))


if __name__ == "__main__":
    main()

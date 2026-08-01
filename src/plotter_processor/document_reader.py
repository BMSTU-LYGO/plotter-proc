from pathlib import Path

from docx import Document

from plotter_processor.models import DocumentText
from plotter_processor.structured_document_reader import (
    read_structured_document,
)

PDF_TEXT_LAYER_ERROR = (
    "PDF does not contain a usable text layer. OCR is not supported in MVP."
)


def read_document(source_path: str | Path) -> DocumentText:
    path = Path(source_path)
    if path.suffix.lower() == ".docx":
        if not path.is_file():
            raise FileNotFoundError(f"Input document does not exist: {path}")
        try:
            legacy_document = Document(path)
        except (OSError, ValueError) as error:
            raise ValueError(f"Cannot read DOCX document: {path}") from error
        return DocumentText([paragraph.text for paragraph in legacy_document.paragraphs], path, [])
    document = read_structured_document(path)
    paragraphs = [
        paragraph
        for page in document.pages
        for element in page.elements
        if hasattr(element, "paragraphs")
        for paragraph in element.paragraphs
    ]
    text = "\n".join(paragraphs)
    if path.suffix.lower() == ".pdf" and len("".join(text.split())) < 10:
        raise ValueError(PDF_TEXT_LAYER_ERROR)
    return DocumentText(paragraphs, path, list(document.warnings))
